import streamlit as st
from docx import Document
import os
from openai import OpenAI

# --- Page Config ---
st.set_page_config(
    page_title="THE SHEIKH EMPIRE",
    page_icon="🔱",
    layout="wide"
)

# --- OpenAI Setup ---
api_key = os.getenv("OPENAI_API_KEY")
client = None
if api_key:
    client = OpenAI(api_key=api_key)

# --- Custom CSS ---
st.markdown("""
<style>
.stApp {
    background: linear-gradient(to right, #eef2f3, #8e9eab);
}
.title {
    font-size:40px; 
    font-weight:700; 
    color:#1f77b4;
    text-align:center;
    margin-bottom:10px;
}
.subtitle {
    font-size:18px; 
    text-align:center;
    color: #333;
    margin-bottom:40px;
}
div.stButton > button:first-child {
    background-color: #1f77b4;
    color: white;
    border-radius: 10px;
    padding: 12px 24px;
    font-size:16px;
    font-weight:600;
    transition:0.3s;
}
div.stButton > button:first-child:hover {
    background-color: #105d8a;
}
</style>
""", unsafe_allow_html=True)

# --- Header ---
st.markdown('<p class="title">🧠 AI Resume & Profile Builder</p>', unsafe_allow_html=True)
st.markdown('<p class="subtitle">Create polished resumes, cover letters, and LinkedIn bios in seconds 🚀</p>', unsafe_allow_html=True)

# --- Sidebar ---
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/3135/3135715.png", width=120)
    st.markdown("## 🔑 About")
    st.write("AI-powered tool to build resumes, cover letters & LinkedIn bios.")
    st.write("💳 Upgrade to **Pro** for premium templates and unlimited downloads.")
    st.write("---")
    st.write("⚡ Powered by OpenAI")

# --- Input Form ---
col1, col2 = st.columns(2)

with col1:
    full_name = st.text_input("👤 Full Name")
    email = st.text_input("📧 Email")
    phone = st.text_input("📞 Phone Number")

with col2:
    skills = st.text_area("💡 Key Skills (comma-separated)")
    experience = st.text_area("💼 Experience")
    education = st.text_area("🎓 Education")

use_ai = st.checkbox("✨ Use AI to polish and rewrite my documents")

# --- Resume Generator ---
if st.button("🚀 Generate Resume & Cover Letter"):
    if not full_name or not email:
        st.error("⚠️ Please fill at least your Name and Email!")
    else:
        # Raw version
        resume_text = f"""
        {full_name}
        {email} | {phone}

        💡 Skills:
        {skills}

        💼 Experience:
        {experience}

        🎓 Education:
        {education}
        """

        cover_letter = f"""
        Dear Hiring Manager,

        I am excited to apply for a role at your company. With my background in {experience}, 
        and skills in {skills}, I am confident I can contribute effectively.

        Looking forward to your response.

        Regards,  
        {full_name}
        """

        linkedin_bio = f"I am {full_name}, skilled in {skills}, with experience in {experience}. Passionate about growth and opportunities."

        # AI enhancement
        if use_ai and client:
            with st.spinner("✨ Polishing with AI..."):
                try:
                    resume_text = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role": "system", "content": "You are an expert resume writer."},
                                  {"role": "user", "content": f"Rewrite this into a professional resume:\n{resume_text}"}]
                    ).choices[0].message.content

                    cover_letter = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role": "system", "content": "You are an expert career coach."},
                                  {"role": "user", "content": f"Write a strong cover letter based on:\nName: {full_name}\nExperience: {experience}\nSkills: {skills}"}]
                    ).choices[0].message.content

                    linkedin_bio = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role": "system", "content": "You are a LinkedIn branding expert."},
                                  {"role": "user", "content": f"Create a LinkedIn About section for {full_name} with skills {skills} and experience {experience}"}]
                    ).choices[0].message.content
                except Exception as e:
                    st.error(f"AI error: {e}")

        # Show results
        st.subheader("📄 Resume")
        st.text_area("", resume_text, height=250)

        st.subheader("✉️ Cover Letter")
        st.text_area("", cover_letter, height=200)

        st.subheader("🔗 LinkedIn Bio")
        st.text_area("", linkedin_bio, height=150)

        # Save DOCX
        doc = Document()
        doc.add_heading(full_name, 0)
        doc.add_paragraph(f"Email: {email} | Phone: {phone}")
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(skills)
        doc.add_heading("Experience", level=1)
        doc.add_paragraph(experience)
        doc.add_heading("Education", level=1)
        doc.add_paragraph(education)
        doc.save("resume.docx")

        with open("resume.docx", "rb") as f:
            st.download_button("⬇️ Download Resume (DOCX)", f, "resume.docx")

        st.download_button("⬇️ Download Cover Letter", cover_letter, "cover_letter.txt")
        st.download_button("⬇️ Download LinkedIn Bio", linkedin_bio, "linkedin_bio.txt")

